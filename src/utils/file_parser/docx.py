"""Word (DOCX) parser implementation."""

import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List

from .models import Block, ParseResult


async def parse_docx(file_path: str, options: Dict = None) -> ParseResult:
    """Parse Word document.
    
    Args:
        file_path: Path to DOCX file
        options: Parser options
        
    Returns:
        ParseResult
    """
    start_time = time.time()
    options = options or {}
    
    file_id = Path(file_path).stem.split("_")[0]
    filename = Path(file_path).name
    
    try:
        import docx
    except ImportError:
        return ParseResult(
            success=False,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_id=file_id,
            filename=filename,
            error="python-docx not installed"
        )
    
    try:
        doc = docx.Document(file_path)
        
        blocks = []
        
        # Extract paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            # Determine if heading
            level = None
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except (ValueError, AttributeError):
                    level = 1
            
            block_type = "heading" if level else "paragraph"
            
            blocks.append(Block(
                chunk_id=f"docx_{para_idx + 1}",
                type=block_type,
                content=text,
                level=level,
                page=para_idx // 30 + 1,  # Approximate page
                method="python-docx",
                confidence=0.95,
                extracted_at=datetime.now().isoformat()
            ))
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables):
            table_blocks = _extract_table_blocks(table, table_idx, blocks[-1].page if blocks else 1)
            blocks.extend(table_blocks)
        
        # Generate markdown
        markdown = _blocks_to_markdown(blocks)
        
        return ParseResult(
            success=True,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_id=file_id,
            filename=filename,
            markdown=markdown,
            blocks=blocks,
            json={
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            },
            parse_time_ms=int((time.time() - start_time) * 1000)
        )
        
    except Exception as e:
        return ParseResult(
            success=False,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            file_id=file_id,
            filename=filename,
            error=str(e),
            parse_time_ms=int((time.time() - start_time) * 1000)
        )


def _extract_table_blocks(table, table_idx: int, page: int) -> List[Block]:
    """Extract blocks from a table."""
    blocks = []
    
    if not table.rows:
        return blocks
    
    # Get all rows
    rows = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)
    
    if not rows:
        return blocks
    
    # Convert to markdown and JSON
    markdown = _table_to_markdown(rows)
    json_table = _table_to_json(rows)
    
    blocks.append(Block(
        chunk_id=f"docx_table_{table_idx + 1}",
        type="table",
        content="",
        markdown=markdown,
        table_json=json_table,
        page=page,
        method="python-docx",
        confidence=0.95,
        extracted_at=datetime.now().isoformat()
    ))
    
    return blocks


def _table_to_markdown(rows: List[List[str]]) -> str:
    """Convert rows to markdown table."""
    if not rows:
        return ""
    
    max_cols = max(len(row) for row in rows)
    col_widths = [max(len(str(row[i])) if i < len(row) else 0 for row in rows) for i in range(max_cols)]
    
    lines = []
    
    # Header
    header_row = rows[0]
    header = "| " + " | ".join(
        str(header_row[i]).ljust(col_widths[i]) if i < len(header_row) else " " * col_widths[i]
        for i in range(len(col_widths))
    ) + " |"
    lines.append(header)
    
    # Separator
    sep = "| " + " | ".join("-" * w for w in col_widths) + " |"
    lines.append(sep)
    
    # Data rows
    for row in rows[1:]:
        data = "| " + " | ".join(
            str(row[i]).ljust(col_widths[i]) if i < len(row) else " " * col_widths[i]
            for i in range(len(col_widths))
        ) + " |"
        lines.append(data)
    
    return "\n".join(lines)


def _table_to_json(rows: List[List[str]]) -> List[List[str]]:
    """Convert rows to JSON."""
    return [[str(cell) for cell in row] for row in rows]


def _blocks_to_markdown(blocks: List[Block]) -> str:
    """Convert blocks to markdown."""
    md_parts = []
    current_page = None
    
    for block in blocks:
        if block.page and block.page != current_page:
            md_parts.append(f"\n--- Page {block.page} ---\n")
            current_page = block.page
        
        if block.type == "heading":
            level = block.level or 1
            md_parts.append(f"{'#' * level} {block.content}\n")
        elif block.type == "paragraph":
            md_parts.append(f"{block.content}\n")
        elif block.type == "table" and block.markdown:
            md_parts.append(f"\n{block.markdown}\n")
    
    return "\n".join(md_parts)
